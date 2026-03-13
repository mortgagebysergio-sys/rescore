import streamlit as st
import fitz
import re
import tempfile
from dataclasses import dataclass, asdict
from typing import List, Optional
from docx import Document

st.set_page_config(page_title="Mortgage Rapid Rescore Strategy Engine", layout="wide")
st.title("Mortgage Rapid Rescore Strategy Engine")

@dataclass
class Tradeline:
    creditor: str
    account_number: str
    status: str
    balance: float
    past_due: float
    high_credit: float
    credit_limit: float
    account_type: str
    source: str
    dla: str
    reported: str
    opened: str
    remarks: str
    recommendation: str = ""
    score_impact: str = ""
    priority: int = 99
    reason: str = ""

def clean_money(value: str) -> float:
    value = (value or "").replace("$", "").replace(",", "").strip()
    if value in {"", "-", "--/--"}:
        return 0.0
    try:
        return float(value)
    except Exception:
        return 0.0

def extract_text_from_pdf(uploaded_file) -> str:
    data = uploaded_file.read()
    doc = fitz.open(stream=data, filetype="pdf")
    pages = []
    for page in doc:
        pages.append(page.get_text("text"))
    return "\n".join(pages)

def normalize_spaces(text: str) -> str:
    return re.sub(r"[ \t]+", " ", text)

def parse_tradelines(text: str) -> List[Tradeline]:
    """
    Parser tuned for merged mortgage credit reports similar to Birchwood layouts.
    It looks for blocks containing Opened/Reported/Balance/Status patterns.
    """
    text = normalize_spaces(text)
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    tradelines = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Likely creditor line
        creditor_match = re.match(r"^([A-Z0-9'&\.\-\/ ]{4,})$", line)
        if creditor_match:
            creditor = creditor_match.group(1).strip()

            # Ignore headers and noise
            bad_headers = {
                "TRADELINES", "TRADE SUMMARY", "DEROGATORY SUMMARY", "CREDITORS",
                "PUBLIC RECORDS", "SOURCE OF INFORMATION", "MISCELLANEOUS INFORMATION"
            }
            if creditor in bad_headers:
                i += 1
                continue

            block = " ".join(lines[i:i+18])

            status = ""
            if "COLLECTION" in block:
                status = "COLLECTION"
            elif "CHARGE OFF" in block or "CHARGED OFF" in block or "PAID CHGOFF" in block or "PAID CHARGE OFF" in block:
                status = "CHARGE OFF"
            elif "CUR WAS 30" in block:
                status = "LATE"
            elif "AS AGREED" in block:
                status = "AS AGREED"
            elif "PAID" in block:
                status = "PAID"

            if status:
                acct_match = re.search(rf"{re.escape(creditor)}\s+([A-Z0-9]+)", block)
                account_number = acct_match.group(1) if acct_match else ""

                opened_match = re.search(r"Opened\s+([0-9]{2}/[0-9]{2})", block)
                reported_match = re.search(r"Reported\s+([0-9]{2}/[0-9]{2})", block)
                high_credit_match = re.search(r"Hi\. Credit\s+\$?([0-9,]+)", block)
                limit_match = re.search(r"Credit Limit\s+\$?([0-9,]+|-)", block)
                past_due_match = re.search(r"Past Due\s+\$?([0-9,\-]+)", block)
                balance_match = re.search(r"Balance\s+\$?([0-9,]+)", block)
                dla_match = re.search(r"DLA\s+([0-9]{2}/[0-9]{2}|--/--)", block)
                source_match = re.search(r"Source \(B\)\s+([A-Z\/]+)", block)
                account_type_match = re.search(r"(Auto|Installment|Revolving|Open)", block)

                remarks = []
                for keyword in [
                    "FIRST PAYMENT NEVER RECEIVED",
                    "COLLECTION ACCOUNT",
                    "PROFIT AND LOSS WRITEOFF; SECURED",
                    "AUTHORIZED USER",
                    "CHARGED OFF ACCOUNT; FIXED RATE",
                    "PAID CHARGE OFF",
                    "CLOSED",
                ]:
                    if keyword in block:
                        remarks.append(keyword)

                tradelines.append(
                    Tradeline(
                        creditor=creditor,
                        account_number=account_number,
                        status=status,
                        balance=clean_money(balance_match.group(1) if balance_match else "0"),
                        past_due=clean_money(past_due_match.group(1) if past_due_match else "0"),
                        high_credit=clean_money(high_credit_match.group(1) if high_credit_match else "0"),
                        credit_limit=clean_money(limit_match.group(1) if limit_match else "0"),
                        account_type=account_type_match.group(1) if account_type_match else "",
                        source=source_match.group(1) if source_match else "",
                        dla=dla_match.group(1) if dla_match else "",
                        reported=reported_match.group(1) if reported_match else "",
                        opened=opened_match.group(1) if opened_match else "",
                        remarks="; ".join(remarks),
                    )
                )
                i += 10
                continue

        i += 1

    return dedupe_tradelines(tradelines)

def dedupe_tradelines(items: List[Tradeline]) -> List[Tradeline]:
    seen = set()
    result = []
    for item in items:
        key = (
            item.creditor.strip(),
            item.account_number.strip(),
            item.status.strip(),
            int(item.balance)
        )
        if key not in seen:
            seen.add(key)
            result.append(item)
    return result

def estimate_revolving_utilization(text: str):
    """
    Reads summary section if present.
    """
    util_match = re.search(r"REVOLVING CREDIT UTILIZATION\s+([0-9]+)%", text)
    total_rev_balance = None
    total_rev_limit = None

    summary_match = re.search(r"REVOLVING\s+[0-9]+\s+([0-9,]+)\s+([0-9,]+)", text)
    if summary_match:
        total_rev_balance = clean_money(summary_match.group(1))
        total_rev_limit = clean_money(summary_match.group(2))

    return {
        "utilization_percent": int(util_match.group(1)) if util_match else None,
        "revolving_balance": total_rev_balance,
        "revolving_limit": total_rev_limit,
    }

def apply_strategy(tradelines: List[Tradeline], util_data: dict) -> List[Tradeline]:
    for t in tradelines:
        status = t.status.upper()
        remarks = t.remarks.upper()

        if status == "COLLECTION":
            if t.balance <= 1000:
                t.recommendation = "Attempt pay-for-delete first"
                t.score_impact = "High: +10 to +35 per deletion"
                t.priority = 1
                t.reason = "Small collections are often the fastest cleanup items for mortgage prep."
            else:
                t.recommendation = "Dispute for accuracy and negotiate deletion/settlement"
                t.score_impact = "Medium to High: +8 to +30"
                t.priority = 2
                t.reason = "Larger collections may require validation or negotiated resolution."

        elif status == "CHARGE OFF":
            if t.balance > 1000:
                t.recommendation = "Send FCRA Section 623 direct dispute"
                t.score_impact = "Medium: +5 to +25 if corrected or removed"
                t.priority = 3
                t.reason = "Higher-balance charge-offs should usually be challenged for reporting accuracy first."
            else:
                t.recommendation = "Dispute first, then consider settlement if validated"
                t.score_impact = "Low to Medium: +3 to +15"
                t.priority = 4
                t.reason = "Smaller charge-offs can sometimes be settled, but deletion is less predictable."

            if "FIRST PAYMENT NEVER RECEIVED" in remarks:
                t.priority = min(t.priority, 2)
                t.reason += " First-payment-default accounts deserve close review for reporting errors."

            if "PAID CHARGE OFF" in remarks or "PAID CHGOFF" in remarks:
                t.recommendation = "Review for inaccuracies; lower urgency unless reporting is wrong"
                t.score_impact = "Low to Medium: +0 to +10"
                t.priority = 6
                t.reason = "Paid charge-offs may still hurt, but live collections and open utilization issues usually rank ahead."

        elif status == "LATE":
            if "AUTHORIZED USER" in remarks:
                t.recommendation = "Consider removal from authorized user account"
                t.score_impact = "Medium: may prevent inherited late-payment damage"
                t.priority = 5
                t.reason = "Authorized-user lates can sometimes be removed by exiting the account."
            else:
                t.recommendation = "Bring current and document payment history"
                t.score_impact = "Low to Medium"
                t.priority = 7
                t.reason = "Late accounts matter, but collections and charge-offs typically rank higher."

        else:
            t.recommendation = "No immediate derogatory action"
            t.score_impact = "Minimal"
            t.priority = 99
            t.reason = "This account is not a primary rapid-rescore target."

    tradelines.sort(key=lambda x: (x.priority, -x.balance))
    return tradelines

def build_utilization_recommendations(util_data: dict):
    recs = []
    util = util_data.get("utilization_percent")
    bal = util_data.get("revolving_balance")
    lim = util_data.get("revolving_limit")

    if util is None or not lim:
        return recs

    target_30 = round(max(0, bal - (lim * 0.30)), 2)
    target_10 = round(max(0, bal - (lim * 0.10)), 2)

    if util > 30:
        recs.append({
            "priority": 1,
            "action": f"Pay revolving balances down below 30% utilization",
            "details": f"Estimated paydown needed: ${target_30:,.2f}",
            "impact": "High: +10 to +35",
            "reason": f"Current revolving utilization is about {util}%, which is hurting scores."
        })

    if util > 10:
        recs.append({
            "priority": 2,
            "action": f"Pay revolving balances down below 10% utilization",
            "details": f"Additional target paydown: ${target_10:,.2f}",
            "impact": "High after cleanup: +5 to +20",
            "reason": "Very low utilization is often helpful once major derogatories are addressed."
        })

    return recs

def generate_summary_targets(tradelines: List[Tradeline], util_recs: List[dict]):
    collections = [t for t in tradelines if t.status == "COLLECTION"]
    chargeoffs = [t for t in tradelines if t.status == "CHARGE OFF"]
    late_au = [t for t in tradelines if t.status == "LATE" and "AUTHORIZED USER" in t.remarks.upper()]

    summary = []
    if util_recs:
        summary.append("Revolving utilization is a top target.")
    if collections:
        summary.append(f"{len(collections)} collection account(s) detected.")
    if chargeoffs:
        summary.append(f"{len(chargeoffs)} charge-off account(s) detected.")
    if late_au:
        summary.append(f"{len(late_au)} authorized-user late account(s) detected.")

    return summary

def create_strategy_doc(tradelines: List[Tradeline], util_data: dict, util_recs: List[dict]) -> str:
    doc = Document()
    doc.add_heading("Mortgage Rapid Rescore Strategy Report", level=1)

    doc.add_paragraph("This report ranks likely credit actions from highest to lowest priority based on the uploaded mortgage credit report.")
    doc.add_heading("Top Priority Action Plan", level=2)

    for rec in util_recs:
        p = doc.add_paragraph(style="List Number")
        p.add_run(rec["action"]).bold = True
        p.add_run(f" | {rec['details']} | Impact: {rec['impact']}")
        doc.add_paragraph(rec["reason"])

    for t in tradelines:
        if t.priority <= 7:
            p = doc.add_paragraph(style="List Number")
            p.add_run(f"{t.creditor} ({t.account_number})").bold = True
            p.add_run(f" | Status: {t.status} | Balance: ${t.balance:,.2f}")
            doc.add_paragraph(f"Recommendation: {t.recommendation}")
            doc.add_paragraph(f"Estimated impact: {t.score_impact}")
            doc.add_paragraph(f"Why: {t.reason}")

    doc.add_heading("Tradeline Detail", level=2)
    for t in tradelines:
        doc.add_heading(f"{t.creditor} - {t.account_number or 'No account # found'}", level=3)
        doc.add_paragraph(f"Status: {t.status}")
        doc.add_paragraph(f"Balance: ${t.balance:,.2f}")
        doc.add_paragraph(f"Past Due: ${t.past_due:,.2f}")
        doc.add_paragraph(f"High Credit: ${t.high_credit:,.2f}")
        doc.add_paragraph(f"Credit Limit: ${t.credit_limit:,.2f}")
        doc.add_paragraph(f"Account Type: {t.account_type}")
        doc.add_paragraph(f"Opened: {t.opened} | Reported: {t.reported} | DLA: {t.dla}")
        doc.add_paragraph(f"Remarks: {t.remarks or 'None captured'}")
        doc.add_paragraph(f"Recommendation: {t.recommendation}")
        doc.add_paragraph(f"Estimated Impact: {t.score_impact}")
        doc.add_paragraph(f"Priority Rank: {t.priority}")

    doc.add_heading("Notes", level=2)
    doc.add_paragraph(
        "Estimated score impacts are directional only. Actual score changes depend on bureau data, model version, age of derogatories, "
        "utilization changes, and whether an item is corrected, updated, or fully deleted."
    )

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name

uploaded_file = st.file_uploader("Upload mortgage credit report PDF", type=["pdf"])

if uploaded_file:
    text = extract_text_from_pdf(uploaded_file)
    tradelines = parse_tradelines(text)
    util_data = estimate_revolving_utilization(text)
    tradelines = apply_strategy(tradelines, util_data)
    util_recs = build_utilization_recommendations(util_data)
    summary = generate_summary_targets(tradelines, util_recs)

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Credit Summary")
        if summary:
            for item in summary:
                st.write(f"- {item}")
        else:
            st.write("No major strategy items detected.")

        st.subheader("Revolving Utilization")
        st.write(util_data)

    with col2:
        st.subheader("Top Action Items")
        shown = 0
        for rec in util_recs:
            st.markdown(f"**Priority {rec['priority']} — {rec['action']}**")
            st.write(rec["details"])
            st.write(f"Impact: {rec['impact']}")
            st.write(rec["reason"])
            shown += 1

        for t in tradelines:
            if t.priority <= 7 and shown < 12:
                st.markdown(f"**Priority {t.priority} — {t.creditor}**")
                st.write(f"Status: {t.status} | Balance: ${t.balance:,.2f}")
                st.write(f"Recommendation: {t.recommendation}")
                st.write(f"Impact: {t.score_impact}")
                st.write(t.reason)
                shown += 1

    st.subheader("Parsed Tradelines")
    display_rows = [asdict(t) for t in tradelines]
    st.dataframe(display_rows, use_container_width=True)

    if st.button("Generate Rapid Rescore Strategy Report"):
        report_path = create_strategy_doc(tradelines, util_data, util_recs)
        with open(report_path, "rb") as f:
            st.download_button(
                label="Download Strategy Report",
                data=f,
                file_name="mortgage_rapid_rescore_strategy.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
